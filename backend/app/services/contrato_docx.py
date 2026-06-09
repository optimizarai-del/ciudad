"""
Generación de contratos en formato Word (.docx) — editable por el admin
antes de firmar. Después puede subir el .docx actualizado a /archivo y queda
asociado al contrato.

Para los tipos de alquiler (alquiler_vivienda / alquiler_comercial) reproduce
el CONTRATO DE LOCACIÓN con el formato y las cláusulas reales de CIUDAD
NEGOCIOS INMOBILIARIOS (mismo modelo que usa la inmobiliaria), prellenando los
datos de la plataforma: locador/a, locataria/o, inmueble, montos, fechas y
garantes. Los demás tipos (boleto / seña) usan un borrador genérico.
"""
import io
from datetime import date, datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


GOLD = RGBColor(0xB8, 0x89, 0x3A)
DARK = RGBColor(0x0A, 0x0A, 0x0A)
GRAY = RGBColor(0x73, 0x73, 0x73)


TIPO_TITULO = {
    "alquiler_vivienda":  "CONTRATO DE LOCACIÓN DE VIVIENDA",
    "alquiler_comercial": "CONTRATO DE LOCACIÓN COMERCIAL",
    "boleto_compraventa": "BOLETO DE COMPRAVENTA",
    "sena_alquiler":      "RESERVA / SEÑA DE ALQUILER",
}

PERIODICIDAD_PALABRA = {
    1: "Mensualmente", 2: "Bimestralmente", 3: "Trimestralmente",
    4: "Cuatrimestralmente", 6: "Semestralmente", 12: "Anualmente",
}

INDICE_DESC = {
    "icl": "el Índice de Contratos de Locación (ICL)",
    "ipc": "el Índice de Precios al Consumidor (IPC)",
    "fijo": "un porcentaje fijo acordado entre las partes",
    "sin_ajuste": "sin ajuste durante toda la vigencia del contrato",
}

MESES = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
         "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]


# ───────────────────────── helpers de formato ──────────────────────────

def _money(v) -> str:
    try:
        n = float(v or 0)
    except Exception:
        return "—"
    return f"$ {n:,.0f}".replace(",", ".") if n > 0 else "—"


def _to_date(d) -> Optional[date]:
    if not d:
        return None
    if isinstance(d, datetime):
        return d.date()
    if isinstance(d, date):
        return d
    try:
        return date.fromisoformat(str(d))
    except Exception:
        return None


def _fecha(d) -> str:
    dd = _to_date(d)
    if not dd:
        return "____ de _______________ de 20__"
    return f"{dd.day:02d} de {MESES[dd.month]} de {dd.year}"


# Número entero a palabras en español (0 .. 999.999.999). Suficiente para
# montos de alquiler/depósito y para días del mes.
_UNI = ['cero', 'uno', 'dos', 'tres', 'cuatro', 'cinco', 'seis', 'siete',
        'ocho', 'nueve', 'diez', 'once', 'doce', 'trece', 'catorce', 'quince',
        'dieciséis', 'diecisiete', 'dieciocho', 'diecinueve', 'veinte',
        'veintiuno', 'veintidós', 'veintitrés', 'veinticuatro', 'veinticinco',
        'veintiséis', 'veintisiete', 'veintiocho', 'veintinueve']
_DEC = ['', '', '', 'treinta', 'cuarenta', 'cincuenta', 'sesenta', 'setenta',
        'ochenta', 'noventa']
_CEN = ['', 'ciento', 'doscientos', 'trescientos', 'cuatrocientos',
        'quinientos', 'seiscientos', 'setecientos', 'ochocientos',
        'novecientos']


def _hasta_99(n: int) -> str:
    if n < 30:
        return _UNI[n]
    d, u = divmod(n, 10)
    return _DEC[d] + ("" if u == 0 else " y " + _UNI[u])


def _hasta_999(n: int) -> str:
    if n == 100:
        return "cien"
    c, r = divmod(n, 100)
    pre = _CEN[c]
    if r == 0:
        return pre
    return (pre + " " if pre else "") + _hasta_99(r)


def _entero_a_letras(n) -> str:
    try:
        n = int(round(float(n)))
    except Exception:
        return ""
    if n == 0:
        return "cero"
    if n < 0:
        return "menos " + _entero_a_letras(-n)
    millones, resto = divmod(n, 1_000_000)
    miles, cien = divmod(resto, 1000)
    partes = []
    if millones:
        partes.append("un millón" if millones == 1 else _hasta_999(millones) + " millones")
    if miles:
        partes.append("mil" if miles == 1 else _hasta_999(miles) + " mil")
    if cien:
        partes.append(_hasta_999(cien))
    return " ".join(partes)


def _monto_en_palabras(v) -> str:
    """'PESOS SEISCIENTOS MIL ($ 600.000)' a partir de 600000."""
    letras = _entero_a_letras(v).upper()
    return f"PESOS {letras} ({_money(v)})" if letras else _money(v)


def _anios_contrato(fi, ff) -> Optional[int]:
    di, df = _to_date(fi), _to_date(ff)
    if not di or not df:
        return None
    meses = (df.year - di.year) * 12 + (df.month - di.month)
    if df.day >= di.day:
        meses += 1
    return max(1, round(meses / 12))


def _persona_str(c: Optional[dict]) -> str:
    """Nombre completo + nacionalidad + documento, para la fórmula de partes."""
    if not c:
        return "[Completar datos de la parte]"
    nombre = (c.get("razon_social") or
              f"{c.get('nombre', '') or ''} {c.get('apellido', '') or ''}".strip())
    bits = [nombre or "[Nombre]"]
    if c.get("nacionalidad"):
        bits.append(c["nacionalidad"])
    tdoc = c.get("tipo_documento") or "DNI"
    if c.get("documento"):
        bits.append(f"{tdoc} N° {c['documento']}")
    return ", ".join(bits)


def _nombre_persona(c: Optional[dict]) -> str:
    if not c:
        return "__________________"
    return (c.get("razon_social") or
            f"{c.get('nombre', '') or ''} {c.get('apellido', '') or ''}".strip()
            or "__________________")


def _direccion_str(p: Optional[dict]) -> str:
    if not p:
        return "[Sin propiedad]"
    bits = [p.get("direccion") or ""]
    if p.get("ciudad"):
        bits.append(p["ciudad"])
    if p.get("provincia"):
        bits.append(p["provincia"])
    return ", ".join([b for b in bits if b])


def _tipo_inmueble(p: Optional[dict]) -> str:
    t = (p or {}).get("tipo") or "inmueble"
    return {
        "departamento": "Departamento", "casa": "Casa", "local": "Local comercial",
        "oficina": "Oficina", "galpon": "Galpón", "campo": "Campo",
    }.get(str(t), str(t).capitalize() or "Inmueble")


# ───────────────────────── primitivas docx ──────────────────────────

def _setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _header_footer(doc: Document):
    """Membrete CIUDAD repetido en cada página, como el contrato original."""
    linea = ("Avda. URUGUAY Nº 268 - SANTA ROSA (L.P.) - 6.300\n"
             "T.E.: 411074 - Móvil: 15-535559 - E-Mail: ciudadinmob@cpenet.com.ar")
    for sec in doc.sections:
        for cont, align in ((sec.header, WD_ALIGN_PARAGRAPH.CENTER),
                            (sec.footer, WD_ALIGN_PARAGRAPH.CENTER)):
            p = cont.paragraphs[0]
            p.alignment = align
            run = p.add_run(linea)
            run.font.size = Pt(8)
            run.font.color.rgb = GRAY


def _h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(8)


def _eyebrow(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = GOLD
    p.paragraph_format.space_after = Pt(2)


def _p(doc, text, justify=True, bold=False, italic=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def _clausula(doc, etiqueta: str, cuerpo: str):
    """Párrafo justificado con la etiqueta de cláusula en negrita/subrayado
    (ej 'PRIMERA a):') seguida del texto."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(etiqueta + " ")
    r.bold = True
    r.underline = True
    r.font.size = Pt(11)
    r2 = p.add_run(cuerpo)
    r2.font.size = Pt(11)
    return p


def _datos_tabla(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (lbl, val) in enumerate(rows):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        c0.width = Cm(5); c1.width = Cm(11)
        r0 = c0.paragraphs[0].add_run(lbl)
        r0.font.size = Pt(10); r0.font.color.rgb = GRAY
        r1 = c1.paragraphs[0].add_run(val or "—")
        r1.font.size = Pt(10); r1.bold = True


# ───────────────────────── contrato de locación CIUDAD ──────────────────────────

def _generar_docx_locacion(*, contrato, propiedad, propietario, inquilino,
                           garantes) -> bytes:
    doc = Document()
    _setup_styles(doc)
    _header_footer(doc)

    comercial = (contrato.get("tipo") == "alquiler_comercial")

    _eyebrow(doc, "CiUDAD")
    _eyebrow(doc, "NEGOCIOS INMOBILIARIOS")
    _h1(doc, "CONTRATO DE LOCACIÓN")

    # ── Encabezado de partes ──
    locador = _persona_str(propietario)
    locataria = _persona_str(inquilino)
    _p(doc,
       f"Entre {locador}, por una parte, en adelante LA LOCADORA, y "
       f"{locataria}, por la otra parte, en adelante LA LOCATARIA, convienen "
       f"en celebrar el siguiente contrato de LOCACIÓN, sujeto a las "
       f"siguientes cláusulas y condiciones:")

    # ── PRIMERA a) — Inmueble y comodidades ──
    tipo_inm = _tipo_inmueble(propiedad)
    direccion = _direccion_str(propiedad)
    desc = (propiedad or {}).get("descripcion") or \
        "[describir comodidades: cantidad de dormitorios, ambientes, cocina, baño, etc.]"
    inventario = (contrato.get("inventario") or "").strip()
    bienes = (f" Los bienes muebles que quedan incorporados en la propiedad "
              f"locada son: {inventario}." if inventario else
              " Los bienes muebles que quedan incorporados en la propiedad "
              "locada constan en archivo de fotos/video.")
    _clausula(doc, "PRIMERA a):",
              f"La parte LOCADORA entrega un inmueble en locación a la parte "
              f"LOCATARIA, quien declara haber visto el mismo y reconocer el "
              f"buen estado en que se encuentra, aceptando en este acto a su "
              f"entera conformidad. RECIBE un {tipo_inm}, ubicado en "
              f"{direccion}, con las siguientes comodidades: {desc}.{bienes}")

    # ── PRIMERA b) — Plazo ──
    anios = _anios_contrato(contrato.get("fecha_inicio"), contrato.get("fecha_fin"))
    if anios:
        plazo = f"{_entero_a_letras(anios).upper()} ({anios}) {'AÑO' if anios == 1 else 'AÑOS'}"
    else:
        plazo = "______ ( ) AÑO(S)"
    _clausula(doc, "PRIMERA b):",
              f"El plazo de esta Locación se fija de común acuerdo por el "
              f"término de {plazo} a partir del {_fecha(contrato.get('fecha_inicio'))} "
              f"con vencimiento el día {_fecha(contrato.get('fecha_fin'))} "
              f"inclusive. La parte LOCATARIA podrá solicitar la rescisión del "
              f"contrato luego de haber transcurrido los 6 primeros meses de "
              f"locación, avisando con una antelación de UN mes (30 días) al "
              f"Locador y/o representante. En tal caso, el LOCATARIO quedará "
              f"expuesto a la indemnización de un mes del alquiler en curso de "
              f"la relación contractual.")

    # ── SEGUNDA — Precio y ajuste ──
    monto = contrato.get("monto_inicial") or 0
    period = int(contrato.get("periodicidad_meses") or 3)
    period_palabra = PERIODICIDAD_PALABRA.get(period, f"cada {period} meses")
    period_num = _entero_a_letras(period).upper()
    indice = (contrato.get("indice_ajuste") or "icl")
    indice_txt = INDICE_DESC.get(indice, "el índice acordado por las partes")
    dia_ini = contrato.get("dia_inicio_pago") or 1
    dia_ven = contrato.get("dia_vencimiento_pago") or 7
    _clausula(doc, "SEGUNDA:",
              f"Queda convenido entre las partes que la parte Locataria "
              f"abonará un alquiler mensual de {_monto_en_palabras(monto)}, más "
              f"tasa municipal y gastos comunes, durante los primeros "
              f"{period_num} ({period}) meses de vigencia del contrato. Para "
              f"los restantes meses de la relación contractual, las partes "
              f"libremente acuerdan corregirlo {period_palabra}, para evitar el "
              f"desequilibrio de las prestaciones recíprocas que genera la "
              f"desvalorización del signo monetario por inflación durante el "
              f"transcurso del contrato. Dicha actualización se efectuará según "
              f"{indice_txt}, publicado mensualmente por el Banco Central de la "
              f"República Argentina, o bien por acuerdo de partes. El pago del "
              f"alquiler será abonado del {dia_ini}° al {dia_ven} de cada mes, "
              f"en el domicilio de la Inmobiliaria interviniente — CIUDAD "
              f"NEGOCIOS INMOBILIARIOS —, con domicilio en Av. Uruguay N° 268, "
              f"de la Ciudad de Santa Rosa, La Pampa.")

    # ── TERCERA — Mora ──
    mora = contrato.get("mora_diaria_porc") or 1
    _clausula(doc, "TERCERA:",
              f"El importe de la locación deberá abonarse por adelantado del "
              f"día {dia_ini} al {dia_ven} de cada mes, y en caso de atraso se "
              f"establece al Locatario una multa diaria del {mora}% del canon "
              f"mensual a abonar por cada día de demora. La Locadora podrá no "
              f"recibir el alquiler si no es satisfecho simultáneamente el "
              f"importe de dicha multa.")

    # ── CUARTA — Destino ──
    destino = ("uso exclusivo comercial, conforme al rubro declarado por la "
               "LOCATARIA" if comercial else
               "uso exclusivo de vivienda familiar, manifestando la LOCATARIA "
               "bajo juramento ocupar la propiedad en forma personal "
               "(habitacional)")
    _clausula(doc, "CUARTA:",
              f"Queda expresamente convenido que el inmueble en cuestión será "
              f"destinado para {destino}, no pudiendo cambiar dicho destino "
              f"bajo ningún motivo, ni parcial ni totalmente. La LOCATARIA "
              f"deberá respetar las normas de convivencia, de buenas "
              f"costumbres, ordenanzas y reglamentaciones propias del edificio "
              f"y demás referidas al uso de la propiedad, y permitir al LOCADOR "
              f"y/o su representante inspeccionar el inmueble previo aviso.")

    # ── QUINTA a DECIMA TERCERA — cláusulas estándar ──
    _clausula(doc, "QUINTA:",
              "La Locadora entrega al ocupante la unidad con todos los "
              "artefactos e instalaciones en general en buen estado, quedando "
              "a cargo de la Locataria los gastos por arreglos que sean "
              "necesarios. El inmueble se entrega pintado con pintura látex de "
              "primera calidad color blanco, debiendo la parte locataria "
              "entregarlo en igual condición.")
    _clausula(doc, "SEXTA:",
              "La parte ocupante deberá constituir, durante toda la vigencia "
              "del presente contrato y hasta la restitución del inmueble, un "
              "Seguro contra incendio, responsabilidad civil y robo, quedando "
              "la parte locadora exenta de tal responsabilidad y obligación.")
    _clausula(doc, "SÉPTIMA:",
              "Queda prohibido: a) ceder la locación, prestar o subarrendar "
              "total o parcialmente, ni provisoria ni permanentemente; b) "
              "realizar cualquier tipo de mejoras o arreglos sin la expresa "
              "autorización escrita de La Locadora; c) realizar cualquier acto "
              "o hecho que pudiera causar daño, molestias y/o perjuicios al "
              "edificio, a los vecinos y/o a terceros. La parte Locataria "
              "deberá respetar las normas de convivencia y los reglamentos de "
              "copropiedad (horarios de descanso, ruidos molestos, etc.).")
    _clausula(doc, "OCTAVA:",
              "La Locadora no será responsable por los daños y perjuicios que "
              "se ocasionaren a los moradores, en sus personas o bienes, desde "
              "que concede la tenencia efectiva del inmueble. En casos de "
              "siniestros, destrucción total o parcial del bien locado o "
              "inundaciones por roturas de cañerías, La Locadora no será "
              "responsable por los daños que pudieran sufrir el grupo familiar, "
              "sus bienes o terceros, ni por accidentes, incendios, derrumbes, "
              "granizos u otro siniestro.")
    _clausula(doc, "NOVENA:",
              "La LOCATARIA deberá poner a su nombre los medidores de energía "
              "eléctrica, gas, teléfono y cable, los cuales estarán a su cargo "
              "exclusivo, además del pago de los servicios por ella "
              "solicitados.")
    _clausula(doc, "DÉCIMA:",
              "Queda convenida la mora automática de La Locataria respecto del "
              "pago del canon locativo, la que se producirá por el mero "
              "vencimiento de los plazos acordados o por el mero incumplimiento "
              "de las obligaciones asumidas, conforme al Código Civil y "
              "Comercial de la Nación y al presente contrato.")
    _clausula(doc, "DÉCIMA PRIMERA:",
              "Para el supuesto de que La Locataria no cumpliera con todas y "
              "cada una de sus obligaciones, La Locadora considerará rescindido "
              "el presente contrato de pleno derecho, debiendo La Locataria "
              "restituir el inmueble libre de toda ocupación dentro de los diez "
              "(10) días corridos de serle notificada tal decisión.")
    _clausula(doc, "DÉCIMA SEGUNDA:",
              f"Si al término de este contrato — que opera el día "
              f"{_fecha(contrato.get('fecha_fin'))}, plazo perentorio e "
              f"improrrogable — o en cualquier otra circunstancia en la que la "
              f"Locataria debiera restituir la unidad, no cumpliese tal "
              f"obligación entregándola libre de ocupación y en perfecto estado "
              f"de uso, conservación y pintado, sin perjuicio de continuar "
              f"devengándose el alquiler correspondiente, La Locataria deberá "
              f"abonar a La Locadora una multa diaria de PESOS "
              f"____________________ ($ __________) por cada día de demora, "
              f"reservándose el derecho a pedir una indemnización. En caso de "
              f"consignación de llaves, el alquiler regirá hasta que el Locador "
              f"tome posesión real y efectiva de la propiedad.")
    _clausula(doc, "DÉCIMA TERCERA:",
              "Queda expresamente facultada La Locadora para que, por sí o por "
              "representante, visite el inmueble alquilado, previa notificación "
              "a la ocupante.")

    # ── DECIMA CUARTA — Garantes ──
    garantes = garantes or []
    if garantes:
        intro = ("Se constituyen en fiadores, lisos, llanos y principales "
                 "pagadores de los alquileres que se devenguen en dicho "
                 "inmueble, renunciando a los beneficios de excusión y "
                 "división, las siguientes personas: ")
        partes_g = []
        for g in garantes:
            nombre = _nombre_persona(g)
            bits = [nombre]
            if g.get("nacionalidad"):
                bits.append(g["nacionalidad"])
            tdoc = g.get("tipo_documento") or "DNI"
            if g.get("documento"):
                bits.append(f"{tdoc} N° {g['documento']}")
            if g.get("domicilio"):
                bits.append(f"con domicilio en {g['domicilio']}")
            if g.get("telefono"):
                bits.append(f"Teléfono: {g['telefono']}")
            partes_g.append(", ".join(bits))
        cuerpo_g = (intro + "; ".join(partes_g) +
                    ". Dicha garantía se extiende a todas y cada una de las "
                    "cláusulas del presente contrato, hasta la efectiva "
                    "restitución del inmueble, afectando todos los bienes "
                    "muebles e inmuebles de los garantes.")
    else:
        cuerpo_g = ("Se constituye(n) en fiador(es), liso(s), llano(s) y "
                    "principal(es) pagador(es), renunciando a los beneficios de "
                    "excusión y división: ____________________________________ "
                    "(nombre, nacionalidad, documento, domicilio y teléfono). "
                    "La garantía se extiende a todas las cláusulas del presente "
                    "contrato hasta la efectiva restitución del inmueble.")
    _clausula(doc, "DÉCIMA CUARTA:", cuerpo_g)

    # ── DECIMA QUINTA — Depósito ──
    deposito = contrato.get("deposito") or 0
    if deposito:
        dep_txt = (f"La parte Locataria entrega en este acto a la parte "
                   f"Locadora la suma de {_monto_en_palabras(deposito)}, en "
                   f"carácter de depósito de garantía")
    else:
        dep_txt = ("La parte Locataria entrega en este acto a la parte "
                   "Locadora la suma de PESOS ____________________ "
                   "($ __________), en carácter de depósito de garantía")
    _clausula(doc, "DÉCIMA QUINTA:",
              f"{dep_txt} por la buena conservación de la unidad y el fiel "
              f"cumplimiento de las obligaciones contraídas, sirviendo el "
              f"presente de suficiente y eficaz recibo. Dicha suma no devenga "
              f"interés alguno, no puede ser aplicada al pago de los alquileres "
              f"y será devuelta a la Locataria una vez finalizada la relación "
              f"contractual, previa constatación y revisión del inmueble. Para "
              f"el caso de que la Locataria no cumpliere con estas "
              f"obligaciones, se le retendrá dicho depósito, más las acciones "
              f"que pudieran corresponder.")

    # ── DECIMA SEXTA — Vía judicial ──
    _clausula(doc, "DÉCIMA SEXTA:",
              "La Locataria autoriza expresamente a La Locadora a solicitar por "
              "vía judicial la desocupación inmediata ante cualquier "
              "incumplimiento, conforme la legislación vigente. La Locadora se "
              "reserva el derecho de exigir el reemplazo de la garantía y/o "
              "exigir una nueva en caso de insolvencia, desaparición o "
              "fallecimiento de los garantes; el Locatario deberá presentar la "
              "nueva garantía dentro de los treinta (30) días de notificado, "
              "bajo apercibimiento de tener por rescindido el contrato como si "
              "se tratara de plazo vencido.")

    # ── DECIMA SEPTIMA — Domicilios y cierre ──
    ciudad_firma = (propiedad or {}).get("ciudad") or "SANTA ROSA (LA PAMPA)"
    hoy = date.today()
    _clausula(doc, "DÉCIMA SÉPTIMA:",
              f"Para todos los efectos legales del presente y de cualquier "
              f"notificación que deba realizarse, las partes constituyen "
              f"domicilios especiales: La Locadora en el domicilio indicado al "
              f"inicio, la Locataria en el domicilio locado, y los garantes en "
              f"los domicilios denunciados. Conforme las partes con todas las "
              f"especificaciones precedentes, se firman tres ejemplares de un "
              f"mismo tenor y a un solo efecto en la Ciudad de {ciudad_firma}, "
              f"a los {_entero_a_letras(hoy.day)} ({hoy.day}) días del mes de "
              f"{MESES[hoy.month]} del año {hoy.year}.")

    if contrato.get("notas"):
        _p(doc, "")
        _p(doc, f"Observaciones: {contrato['notas']}", italic=True)

    # ── Firmas ──
    _p(doc, "")
    _p(doc, "")
    t = doc.add_table(rows=2, cols=3)
    for j in range(3):
        t.cell(0, j).paragraphs[0].add_run("____________________").font.size = Pt(10)
    etiquetas = ["LOCATARIA", "LOCADOR", "GARANTES"]
    for j, et in enumerate(etiquetas):
        run = t.cell(1, j).paragraphs[0].add_run(et)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ───────────────────────── borrador genérico (boleto / seña) ──────────────────────────

def _generar_docx_generico(*, contrato, propiedad, propietario, inquilino) -> bytes:
    doc = Document()
    _setup_styles(doc)

    tipo = contrato.get("tipo") or "alquiler_vivienda"
    titulo = TIPO_TITULO.get(tipo, "CONTRATO")

    _eyebrow(doc, "CiUDAD · NEGOCIOS INMOBILIARIOS")
    _h1(doc, titulo)

    _p(doc, "")
    _p(doc,
       f"En la ciudad de _________________, a los {_fecha(contrato.get('fecha_inicio'))}, "
       f"entre las partes que a continuación se identifican, se celebra el presente "
       f"contrato bajo las cláusulas y condiciones que se detallan.")

    rol_a = "LOCADOR" if tipo.startswith("alquiler") else "VENDEDOR"
    rol_b = "LOCATARIO" if tipo.startswith("alquiler") else "COMPRADOR"
    _datos_tabla(doc, [
        (rol_a, _persona_str(propietario)),
        (rol_b, _persona_str(inquilino)),
    ])

    obj_txt = {
        "boleto_compraventa":
            f"El VENDEDOR vende y el COMPRADOR adquiere el inmueble ubicado en {_direccion_str(propiedad)}.",
        "sena_alquiler":
            f"El presente instrumenta la reserva del inmueble ubicado en {_direccion_str(propiedad)} "
            f"con destino a alquiler, sujeto a la firma del contrato definitivo.",
    }.get(tipo, f"Inmueble ubicado en {_direccion_str(propiedad)}.")
    _p(doc, "")
    _p(doc, obj_txt)

    _datos_tabla(doc, [
        ("Inicio", _fecha(contrato.get("fecha_inicio"))),
        ("Vencimiento", _fecha(contrato.get("fecha_fin"))),
        ("Monto inicial", _money(contrato.get("monto_inicial"))),
        ("Depósito / Seña", _money(contrato.get("deposito"))),
        ("Comisión inmobiliaria", f"{contrato.get('comision_porc') or 0} %"),
    ])

    if contrato.get("notas"):
        _p(doc, "")
        _p(doc, f"Observaciones: {contrato['notas']}", italic=True)

    _p(doc, "")
    _p(doc, "")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "____________________________"
    t.cell(0, 1).text = "____________________________"
    t.cell(1, 0).text = rol_a
    t.cell(1, 1).text = rol_b
    for cell in (t.cell(1, 0), t.cell(1, 1)):
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = GRAY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_docx(*, contrato: dict, propiedad: dict, propietario: Optional[dict],
                 inquilino: Optional[dict], garantes: Optional[list] = None) -> bytes:
    """Devuelve el .docx en bytes. Para alquileres usa el modelo de locación
    CIUDAD; para boleto/seña, un borrador genérico."""
    tipo = (contrato or {}).get("tipo") or "alquiler_vivienda"
    if tipo in ("alquiler_vivienda", "alquiler_comercial"):
        return _generar_docx_locacion(
            contrato=contrato, propiedad=propiedad, propietario=propietario,
            inquilino=inquilino, garantes=garantes or [],
        )
    return _generar_docx_generico(
        contrato=contrato, propiedad=propiedad, propietario=propietario,
        inquilino=inquilino,
    )
