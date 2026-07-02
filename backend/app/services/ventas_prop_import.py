"""
Carga de propiedad de VENTA desde archivo con IA (Anthropic Claude) — Fase 4.4.

Subís la ficha del portal, el aviso, el plano municipal o el PDF de la propiedad
y Claude extrae los datos estructurados para cargarla al catálogo de ventas en
segundos (en lugar de tipear todo a mano). Cierra la brecha "carga desde plano".

Flujo:
1. El frontend manda el archivo (PDF / imagen JPG-PNG / DOCX / TXT).
2. Se pasa a Claude con un prompt que pide el esquema de VentasPropiedad.
3. `/preview` devuelve el JSON sin guardar → el operador revisa el form.
4. `/confirmar` recibe el JSON revisado y crea la VentasPropiedad, con geo +
   matching (para que aparezca en el mapa y en Matches).

Requiere ANTHROPIC_API_KEY en el entorno. Si no está, se lanza un error claro.
"""
from __future__ import annotations

import base64
import json
import os
import re


SYSTEM_PROMPT = """Sos un asistente que extrae datos estructurados de fichas,
avisos, planos o PDFs de propiedades inmobiliarias ARGENTINAS en VENTA, para
cargarlas en un catálogo. Tu única salida es JSON válido, sin markdown ni texto
adicional.

Esquema de respuesta:
{
  "titulo": <string breve y descriptivo, ej "Casa 3 dorm. c/ patio en Barrio Palermo">,
  "tipo": "casa" | "departamento" | "lote" | "local" | "oficina" | "galpon" | "campo" | "otro",
  "direccion": <string: calle y número si figura> | null,
  "ciudad": <string: localidad> | null,
  "provincia": <string> | null,
  "precio_usd": <number en USD, sin símbolos ni miles> | null,
  "moneda_original": "USD" | "ARS" | null,
  "superficie_m2": <number> | null,
  "dormitorios": <number> | null,
  "banos": <number> | null,
  "antiguedad_anios": <number> | null,
  "descripcion": <string: descripción completa con comodidades/amenities> | null,
  "inmobiliaria": <string: quién publica, si figura> | null,
  "contacto": { "nombre": <string>|null, "telefono": <string>|null, "email": <string>|null } | null
}

Reglas:
- Si no podés determinar un campo, devolvé null (NO inventes datos).
- Es un servicio SOLO para Argentina: la provincia y ciudad deben ser argentinas.
- precio_usd: los inmuebles en venta en Argentina casi siempre se listan en
  dólares. Si el precio está en USD, poné el número en precio_usd y
  moneda_original="USD". Si está SOLO en pesos, dejá precio_usd=null y
  moneda_original="ARS" (no conviertas: no sabés la cotización del día).
- tipo: mapeá PH/monoambiente → "departamento"; terreno → "lote"; galpón/depósito
  → "galpon"; chacra/quinta → "campo"; cochera → "otro".
- Números sin puntos de miles ni símbolos. "USD 145.000" → 145000.
- superficie_m2: usá la superficie cubierta si hay varias; si solo hay total, usá esa.
- descripcion: redactá una descripción completa a partir de las comodidades que
  liste el aviso (dormitorios, ambientes, cocina, patio, cochera, estado, orientación,
  amenities). No inventes comodidades que no figuren.
- titulo: armá un título corto y comercial combinando tipo + característica principal + zona.
"""


def _client():
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY no está configurada en el servidor.")
    from anthropic import Anthropic
    return Anthropic(api_key=api_key)


_IMG_MEDIA = {
    ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
    ".webp": "image/webp", ".gif": "image/gif",
}


def parsear_propiedad_archivo(content: bytes, filename: str) -> dict:
    """Extrae los datos de la propiedad desde un archivo (PDF/imagen/DOCX/TXT)."""
    client = _client()
    fn = (filename or "").lower()
    ext = "." + fn.rsplit(".", 1)[-1] if "." in fn else ""

    if ext == ".pdf":
        return _claude_documento(client, content, "application/pdf")
    if ext in _IMG_MEDIA:
        return _claude_imagen(client, content, _IMG_MEDIA[ext])
    if ext == ".docx":
        texto = _texto_de_docx(content)
        return _claude_texto(client, texto)
    if ext in (".txt", ".text"):
        try:
            texto = content.decode("utf-8")
        except UnicodeDecodeError:
            texto = content.decode("latin-1", errors="ignore")
        return _claude_texto(client, texto)
    # Default: intentar como PDF
    return _claude_documento(client, content, "application/pdf")


def _texto_de_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ValueError("Falta el paquete python-docx en el servidor.")
    from io import BytesIO
    doc = Document(BytesIO(content))
    out = [p.text for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    out.append(cell.text)
    texto = "\n".join(out).strip()
    if not texto:
        raise ValueError("No se pudo extraer texto del archivo.")
    return texto


_MODEL = os.getenv("ANTHROPIC_MODEL_IMPORT", "claude-sonnet-4-5")
_INSTR = ("Extraé los datos estructurados de esta propiedad según el esquema "
          "indicado en las instrucciones. Devolvé SOLO el JSON, sin texto extra.")


def _claude_texto(client, texto: str) -> dict:
    if not (texto or "").strip():
        raise ValueError("No se pudo extraer texto del archivo.")
    texto = texto[:60_000]
    msg = client.messages.create(
        model=_MODEL, max_tokens=2048, system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content":
                   f"{_INSTR}\n\n---INICIO FICHA---\n{texto}\n---FIN FICHA---"}])
    return _parse(msg)


def _claude_documento(client, data: bytes, media_type: str) -> dict:
    b64 = base64.standard_b64encode(data).decode("ascii")
    msg = client.messages.create(
        model=_MODEL, max_tokens=2048, system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": [
            {"type": "document",
             "source": {"type": "base64", "media_type": media_type, "data": b64}},
            {"type": "text", "text": _INSTR},
        ]}])
    return _parse(msg)


def _claude_imagen(client, data: bytes, media_type: str) -> dict:
    b64 = base64.standard_b64encode(data).decode("ascii")
    msg = client.messages.create(
        model=_MODEL, max_tokens=2048, system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": [
            {"type": "image",
             "source": {"type": "base64", "media_type": media_type, "data": b64}},
            {"type": "text", "text": _INSTR},
        ]}])
    return _parse(msg)


_TIPOS_VALIDOS = {"casa", "departamento", "lote", "local", "oficina", "galpon", "campo", "otro"}
_TIPO_MAP = {
    "ph": "departamento", "monoambiente": "departamento", "depto": "departamento",
    "terreno": "lote", "deposito": "galpon", "depósito": "galpon",
    "galpón": "galpon", "chacra": "campo", "quinta": "campo", "cochera": "otro",
}


def _parse(msg) -> dict:
    out = "".join(b.text for b in msg.content if getattr(b, "type", None) == "text").strip()
    if out.startswith("```"):
        out = re.sub(r"^```(?:json)?\s*", "", out)
        out = re.sub(r"\s*```\s*$", "", out)
    try:
        data = json.loads(out)
    except json.JSONDecodeError as e:
        raise ValueError(f"Claude no devolvió JSON parseable: {e}. Respuesta: {out[:400]}")
    # Normalizar tipo al enum
    t = (data.get("tipo") or "otro").strip().lower()
    t = _TIPO_MAP.get(t, t)
    data["tipo"] = t if t in _TIPOS_VALIDOS else "otro"
    for k in ("titulo", "direccion", "ciudad", "provincia", "descripcion", "inmobiliaria"):
        data.setdefault(k, None)
    for k in ("precio_usd", "superficie_m2", "dormitorios", "banos", "antiguedad_anios"):
        data.setdefault(k, None)
    data.setdefault("contacto", None)
    return data
