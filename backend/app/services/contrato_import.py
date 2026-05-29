"""
Importación de contratos desde PDF con IA (Anthropic Claude).

Flujo:
1. El frontend manda el PDF como multipart/form-data.
2. Lo pasamos en base64 a Claude con un prompt que pide extraer:
   - Datos del propietario (locador)
   - Datos del inquilino (locatario)
   - Datos del inmueble (dirección, tipo, etc.)
   - Datos del contrato (tipo, fechas, monto, depósito, comisión, ajuste)
3. Claude devuelve JSON estructurado.
4. El endpoint /preview lo retorna sin guardar; el frontend muestra el form
   prellenado para que el operador revise.
5. /confirmar recibe el JSON revisado y crea las entidades en orden
   (clientes → propiedad → contrato), reutilizando registros existentes
   por DNI o por dirección.

Requiere ANTHROPIC_API_KEY en el entorno.
"""
import base64
import json
import os
import re
from typing import Optional

from sqlalchemy.orm import Session

from app import models
from app.services.workspace import workspace_flag


SYSTEM_PROMPT = """Sos un asistente que extrae datos estructurados de contratos
inmobiliarios argentinos (alquileres y boletos de compraventa). Tu única salida
es JSON válido, sin markdown ni texto adicional.

Esquema de respuesta:
{
  "tipo": "alquiler_vivienda" | "alquiler_comercial" | "boleto_compraventa" | "sena_alquiler",
  "estado": "borrador" | "vigente" | "vencido" | "rescindido" | "reservado",
  "fecha_inicio": "YYYY-MM-DD" | null,
  "fecha_fin": "YYYY-MM-DD" | null,
  "monto_inicial": <number> | null,
  "deposito": <number> | null,
  "indice_ajuste": "ipc" | "icl" | "fijo" | "sin_ajuste",
  "periodicidad_meses": <number, default 6>,
  "porcentaje_fijo": <number> | null,
  "comision_porc": <number> | null,
  "mora_diaria_porc": <number> | null,
  "dia_inicio_pago": <number 1-31> | null,
  "dia_vencimiento_pago": <number 1-31> | null,
  "inventario": <string largo> | null,
  "notas": <string breve resumiendo cláusulas relevantes>,

  "propietarios": [
    {
      "nombre": <string>,
      "apellido": <string> | null,
      "razon_social": <string> | null,
      "documento": <DNI o CUIT con guiones si los hay> | null,
      "tipo_documento": "DNI" | "CUIT" | "CUIL" | "Pasaporte" | "LE" | "LC" | "Otro" | null,
      "nacionalidad": <string> | null,
      "email": <string> | null,
      "telefono": <string> | null,
      "porcentaje": <number 0-100> | null
    }
    // ... uno o más; si en el contrato firma sólo uno, devolvé el array con 1
    // si firman 2+ (matrimonio, hermanos, sociedad de hecho) devolvé todos
  ],
  // Inquilino (PRINCIPAL — único, para retrocompatibilidad). Coincide con
  // el primer elemento de inquilinos[] si esa lista vino.
  "inquilino": {
    "nombre": <string>,
    "apellido": <string> | null,
    "razon_social": <string> | null,
    "documento": <string> | null,
    "tipo_documento": "DNI" | "CUIT" | "CUIL" | "Pasaporte" | "LE" | "LC" | "Otro" | null,
    "nacionalidad": <string> | null,
    "email": <string> | null,
    "telefono": <string> | null
  },
  "inquilinos": [
    // Si el contrato tiene MÁS DE UN inquilino/locatario firmante
    // (matrimonios, hermanos compartiendo el alquiler, sociedades),
    // devolvelos todos acá. Cada item: mismo formato que `inquilino`.
    // El primero se considera principal por defecto.
    // Si hay un solo inquilino, dejá esta lista vacía o ponelo solo en
    // el campo `inquilino` de arriba.
    { ... }
  ],
  "co_firmantes": [
    // Garantes / codeudores / fiadores que firman el contrato.
    // Mismo formato que inquilino. Lista vacía si no hay.
    { ... }
  ],
  "propiedad": {
    "direccion": <string>,
    "ciudad": <string> | null,
    "provincia": <string> | null,
    "tipo": "departamento" | "casa" | "local" | "oficina" | "galpon" | "campo",
    "modalidad": "alquiler" | "venta" | "ambas",
    "superficie_m2": <number> | null,
    "ambientes": <number> | null,
    "descripcion": <string breve> | null,
    "precio_alquiler": <number> | null,
    "expensas": <number> | null,
    "tasa_municipal": <number> | null
  },
  "policies": {
    // Quién paga cada concepto según el contrato. Sirve para pre-marcar
    // los toggles en el modal de cobranza.
    "tasas_municipales_a_cargo": "inquilino" | "propietario" | null,
    "expensas_a_cargo":          "inquilino" | "propietario" | null,
    "impuesto_inmobiliario":     "inquilino" | "propietario" | null,
    "servicios_a_cargo":         "inquilino" | "propietario" | null,
    "seguro_obligatorio":        <bool>,
    "mascotas_permitidas":       <bool> | null
  }
}

Reglas:
- Si no podés determinar un campo, devolvé null (NO inventes).
- Montos en pesos argentinos como número, sin símbolos ni miles. Ejemplo: 350000.
- Las fechas siempre en formato ISO YYYY-MM-DD. Si el contrato dice "1 de enero
  de 2026", devolvé "2026-01-01".
- Si el contrato menciona un ajuste por IPC pero no aclara periodicidad,
  asumí 6 meses. Si dice "trimestral" → 3, "semestral" → 6, "anual" → 12.
- "monto_inicial" es el valor del primer mes del alquiler (o el precio total
  si es boleto de compraventa). Si hay un período fijo inicial al mismo monto
  (ej: "los primeros 3 meses al mismo valor"), usá ese monto.
- Para los nombres: separá nombre y apellido si vienen juntos. Si es una
  empresa, usá `razon_social` y dejá `nombre`/`apellido` vacíos.
- Si encontrás CUIT/CUIL en formato XX-XXXXXXXX-X mantené los guiones.
- Documentos: copialos TAL CUAL aparecen, incluyendo puntos y guiones.
  No los normalices.
- tipo_documento: si dice DNI/D.N.I. → "DNI"; si tiene formato XX-XXXXXXXX-X
  y dice CUIT → "CUIT"; CUIL similar. Si no hay aclaración pero parece DNI
  argentino (7-8 dígitos con o sin puntos), usá "DNI".
- mora_diaria_porc: si el contrato dice "punición diaria del 1%" → 1.0.
- dia_inicio_pago / dia_vencimiento_pago: si dice "del 1° al 7 de cada mes"
  → dia_inicio_pago=1, dia_vencimiento_pago=7.
- Múltiples propietarios: si el contrato firma como locadores 2 o más
  personas (matrimonio, herederos, sociedad), devolvelos a todos en
  `propietarios`. Si el contrato menciona explícitamente porcentajes de
  copropiedad, completalos en `porcentaje`; sino dejá `porcentaje: null`.
- Múltiples inquilinos: si firman como LOCATARIOS 2 o más personas
  (matrimonios, compañeros de departamento, sociedades comerciales),
  devolvelos a todos en `inquilinos`. El primero es el titular principal.
  Si firma uno solo, dejá `inquilinos: []` y completá `inquilino` arriba.
  NO incluir acá a garantes/fiadores — esos van en `co_firmantes`.
- co_firmantes: garantes, fiadores y codeudores van acá. NO son inquilinos
  principales pero hay que cargarlos como Clientes para tener su contacto.
  Si la cláusula dice "los locatarios serán solidariamente responsables"
  pero el firmante es una sola persona, NO inventes co-firmantes.
- inventario: copialo casi completo si está, listando muebles/electrodomésticos.
- policies: en contratos típicos de vivienda, expensas y servicios son del
  inquilino, e impuesto inmobiliario es del propietario. No asumas si no
  está explícito; en ese caso null.
"""


def _texto_de_docx(content: bytes) -> str:
    """Extrae texto de un .docx (Word moderno)."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("Falta el paquete python-docx en el servidor.")
    from io import BytesIO
    doc = Document(BytesIO(content))
    out = []
    for p in doc.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    out.append(cell.text)
    return "\n".join(out).strip()


def _texto_de_doc(content: bytes) -> str:
    """Extrae texto de un .doc (Word binario viejo, OLE Compound File).
    No es perfecto pero Claude maneja bien el ruido."""
    try:
        import olefile
    except ImportError:
        raise ValueError("Falta el paquete olefile en el servidor.")
    from io import BytesIO
    ole = olefile.OleFileIO(BytesIO(content))
    if not ole.exists("WordDocument"):
        return ""
    data = ole.openstream("WordDocument").read()
    txt = data.decode("latin-1", errors="ignore")
    txt = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]+", " ", txt)
    txt = re.sub(r" {2,}", " ", txt)
    return txt.strip()


def parsear_contrato_archivo(content: bytes, filename: str) -> dict:
    """Parsea un contrato en cualquier formato soportado (PDF/DOC/DOCX/TXT).
    Para PDF, envía el archivo crudo a Claude (mejor calidad con tablas/imágenes).
    Para DOC/DOCX/TXT extrae texto plano y lo envía al LLM.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY no está configurada en el servidor.")

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    fn = (filename or "").lower()
    if fn.endswith(".pdf"):
        return _llamar_claude_pdf(client, content)
    if fn.endswith(".docx"):
        texto = _texto_de_docx(content)
    elif fn.endswith(".doc"):
        texto = _texto_de_doc(content)
    elif fn.endswith(".txt") or fn.endswith(".text"):
        try:
            texto = content.decode("utf-8")
        except UnicodeDecodeError:
            texto = content.decode("latin-1", errors="ignore")
    else:
        # Default: tratar como PDF
        return _llamar_claude_pdf(client, content)

    if not texto.strip():
        raise ValueError("No se pudo extraer texto del archivo.")
    return _llamar_claude_texto(client, texto)


def _llamar_claude_texto(client, texto: str) -> dict:
    """Envía el texto del contrato al LLM y parsea la respuesta JSON."""
    # Truncar a algo razonable
    MAX = 80_000
    if len(texto) > MAX:
        texto = texto[:MAX] + "\n\n[...texto truncado...]"
    msg = client.messages.create(
        model=os.getenv("ANTHROPIC_MODEL_IMPORT", "claude-sonnet-4-5"),
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[{
            "role": "user",
            "content": (
                "Acá tenés el texto del contrato. Devolvé el JSON estructurado:\n\n"
                f"---INICIO CONTRATO---\n{texto}\n---FIN CONTRATO---"
            ),
        }],
    )
    return _parse_response(msg)


def _llamar_claude_pdf(client, pdf_bytes: bytes) -> dict:
    """Envía el PDF crudo a Claude (que sabe leer PDFs nativamente)."""
    b64 = base64.standard_b64encode(pdf_bytes).decode("ascii")
    msg = client.messages.create(
        model=os.getenv("ANTHROPIC_MODEL_IMPORT", "claude-sonnet-4-5"),
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": b64,
                        },
                    },
                    {
                        "type": "text",
                        "text": (
                            "Extraé los datos estructurados de este contrato "
                            "inmobiliario según el esquema indicado en las "
                            "instrucciones. Devolvé SOLO el JSON, sin texto extra."
                        ),
                    },
                ],
            }
        ],
    )
    return _parse_response(msg)


def _parse_response(msg) -> dict:
    """Limpia y parsea la respuesta JSON del LLM (con todas las quirks)."""

    # Concatenar bloques de texto
    out = "".join(b.text for b in msg.content if getattr(b, "type", None) == "text")

    # Claude a veces envuelve en ```json … ``` — limpiar
    out = out.strip()
    if out.startswith("```"):
        out = re.sub(r"^```(?:json)?\s*", "", out)
        out = re.sub(r"\s*```\s*$", "", out)

    try:
        data = json.loads(out)
    except json.JSONDecodeError as e:
        raise ValueError(f"Claude no devolvió JSON parseable: {e}. Respuesta: {out[:500]}")

    # Compatibilidad: si el modelo devolvió el formato viejo `propietario: {}`
    # lo convertimos a la lista nueva `propietarios: [{}]`.
    if "propietario" in data and "propietarios" not in data:
        single = data.pop("propietario")
        if single and (single.get("nombre") or single.get("razon_social")):
            data["propietarios"] = [single]
        else:
            data["propietarios"] = []
    # Asegurar las claves obligatorias
    data.setdefault("propietarios", [])
    data.setdefault("inquilino", {})
    data.setdefault("co_firmantes", [])
    data.setdefault("propiedad", {})
    data.setdefault("policies", {})
    return data


# Compat: alias de la función vieja para que código existente siga funcionando
def parsear_contrato_pdf(pdf_bytes: bytes, filename: str = "contrato.pdf") -> dict:
    """Alias retrocompatible — usa parsear_contrato_archivo."""
    return parsear_contrato_archivo(pdf_bytes, filename)


def _buscar_cliente_existente(
    db: Session, datos: dict, rol: str, is_demo: bool
) -> Optional[models.Cliente]:
    """Busca por documento (preferido) o por nombre+apellido si no hay doc."""
    doc = (datos.get("documento") or "").strip()
    q = db.query(models.Cliente).filter(models.Cliente.is_demo == is_demo)
    if doc:
        match = q.filter(models.Cliente.documento == doc).first()
        if match:
            return match
    razon = (datos.get("razon_social") or "").strip()
    if razon:
        match = q.filter(models.Cliente.razon_social.ilike(razon)).first()
        if match:
            return match
    nombre = (datos.get("nombre") or "").strip()
    apellido = (datos.get("apellido") or "").strip()
    if nombre and apellido:
        match = q.filter(
            models.Cliente.nombre.ilike(nombre),
            models.Cliente.apellido.ilike(apellido),
        ).first()
        if match:
            return match
    return None


def _buscar_propiedad_existente(
    db: Session, datos: dict, is_demo: bool
) -> Optional[models.Propiedad]:
    dir_norm = (datos.get("direccion") or "").strip()
    if not dir_norm:
        return None
    q = db.query(models.Propiedad).filter(models.Propiedad.is_demo == is_demo)
    # Igualdad case-insensitive
    match = q.filter(models.Propiedad.direccion.ilike(dir_norm)).first()
    return match


def crear_desde_parsed(db: Session, datos: dict, user) -> dict:
    """Crea/reutiliza propietario, inquilino, propiedad y contrato a partir
    del JSON extraído. Devuelve un resumen de lo que hizo (ids + flags
    'reutilizado' por entidad).

    No hace commit final hasta tener todo armado para ser transaccional.
    """
    is_demo = workspace_flag(user)
    resumen = {
        "propietarios": [],   # lista de {id, nombre, reutilizado, porcentaje}
        "inquilino":   {"id": None, "reutilizado": False},
        "propiedad":   {"id": None, "reutilizado": False},
        "contrato":    {"id": None, "codigo": None},
    }

    # 1. Propietarios — lista de uno o más
    propietarios_data = datos.get("propietarios") or []
    # Compat: si vino el formato viejo {propietario: {}}, convertirlo
    if not propietarios_data and datos.get("propietario"):
        propietarios_data = [datos["propietario"]]
    if not propietarios_data:
        raise ValueError("Falta al menos un propietario para poder importar.")

    propietarios_creados = []  # tuplas (cliente, porcentaje, reutilizado)
    for prop_data in propietarios_data:
        if not (prop_data.get("nombre") or prop_data.get("razon_social")):
            continue
        existing = _buscar_cliente_existente(db, prop_data, "propietario", is_demo)
        if existing:
            # Si el cliente existe pero no tiene tipo_documento o nacionalidad,
            # los completamos con lo que vino del contrato (sin pisar lo existente).
            if not existing.tipo_documento and prop_data.get("tipo_documento"):
                existing.tipo_documento = prop_data["tipo_documento"]
            if not existing.nacionalidad and prop_data.get("nacionalidad"):
                existing.nacionalidad = prop_data["nacionalidad"]
            propietarios_creados.append((existing, prop_data.get("porcentaje"), True))
        else:
            nuevo = models.Cliente(
                nombre=prop_data.get("nombre") or (prop_data.get("razon_social") or "Propietario"),
                apellido=prop_data.get("apellido"),
                razon_social=prop_data.get("razon_social"),
                documento=prop_data.get("documento"),
                tipo_documento=prop_data.get("tipo_documento"),
                nacionalidad=prop_data.get("nacionalidad"),
                email=prop_data.get("email"),
                telefono=prop_data.get("telefono"),
                rol=models.ClienteRol.propietario,
                notas="[IMPORTADO desde contrato]",
                is_demo=is_demo,
            )
            db.add(nuevo); db.flush()
            propietarios_creados.append((nuevo, prop_data.get("porcentaje"), False))

    if not propietarios_creados:
        raise ValueError("Ningún propietario válido en el PDF (faltan nombres).")

    # El principal será el primero
    propietario = propietarios_creados[0][0]
    for cli, porc, reu in propietarios_creados:
        resumen["propietarios"].append({
            "id": cli.id,
            "nombre": (cli.razon_social or
                       f"{cli.nombre} {cli.apellido or ''}".strip()),
            "porcentaje": porc,
            "reutilizado": reu,
        })

    # 2. Inquilinos — uno o varios firmantes del contrato.
    # Por compat con extractos viejos, aceptamos tanto `inquilinos` (lista
    # nueva) como `inquilino` (objeto único). Si vienen los dos, unificamos:
    # el objeto único pasa a ser el primero de la lista.
    inquilinos_data: list[dict] = list(datos.get("inquilinos") or [])
    inq_single = datos.get("inquilino") or {}
    if inq_single and (inq_single.get("nombre") or inq_single.get("razon_social") or inq_single.get("documento")):
        # Si la lista nueva ya incluye al principal por documento, no duplicar
        doc_principal = (inq_single.get("documento") or "").strip()
        ya_listado = any(
            (i.get("documento") or "").strip() == doc_principal and doc_principal
            for i in inquilinos_data
        )
        if not ya_listado:
            inquilinos_data.insert(0, inq_single)

    inquilinos_creados: list = []  # tuplas (cliente, reutilizado, es_principal)
    for idx, inq in enumerate(inquilinos_data):
        if not (inq.get("nombre") or inq.get("razon_social") or inq.get("documento")):
            continue
        existing = _buscar_cliente_existente(db, inq, "inquilino", is_demo)
        if existing:
            if not existing.tipo_documento and inq.get("tipo_documento"):
                existing.tipo_documento = inq["tipo_documento"]
            if not existing.nacionalidad and inq.get("nacionalidad"):
                existing.nacionalidad = inq["nacionalidad"]
            inquilinos_creados.append((existing, True, idx == 0))
        else:
            nuevo = models.Cliente(
                nombre=inq.get("nombre") or (inq.get("razon_social") or "Inquilino"),
                apellido=inq.get("apellido"),
                razon_social=inq.get("razon_social"),
                documento=inq.get("documento"),
                tipo_documento=inq.get("tipo_documento"),
                nacionalidad=inq.get("nacionalidad"),
                email=inq.get("email"),
                telefono=inq.get("telefono"),
                rol=models.ClienteRol.inquilino,
                notas="[IMPORTADO desde contrato]",
                is_demo=is_demo,
            )
            db.add(nuevo); db.flush()
            inquilinos_creados.append((nuevo, False, idx == 0))

    # Inquilino principal = el primero (o None si no hubo ninguno)
    inquilino = inquilinos_creados[0][0] if inquilinos_creados else None

    # Resumen retrocompatible
    resumen["inquilinos"] = [
        {
            "id":            cli.id,
            "nombre":        (cli.razon_social or
                              f"{cli.nombre} {cli.apellido or ''}".strip()),
            "documento":     cli.documento,
            "reutilizado":   reu,
            "es_principal":  principal,
        }
        for (cli, reu, principal) in inquilinos_creados
    ]
    if inquilino:
        resumen["inquilino"] = {"id": inquilino.id, "reutilizado": inquilinos_creados[0][1]}

    # 2bis. Co-firmantes (garantes, codeudores, fiadores).
    # Se crean como Clientes con rol=inquilino y notas indicando que son
    # garantes — para tener su contacto cargado en la base. Si querés
    # vincularlos formalmente al contrato en una tabla pivote, agregalo
    # como mejora futura (por ahora alcanza con que estén cargados).
    cofirmantes_data = datos.get("co_firmantes") or []
    cofirmantes_creados = []
    for cf in cofirmantes_data:
        if not (cf.get("nombre") or cf.get("razon_social")):
            continue
        ex = _buscar_cliente_existente(db, cf, "inquilino", is_demo)
        if ex:
            cofirmantes_creados.append({"id": ex.id, "reutilizado": True,
                                         "nombre": (ex.razon_social or
                                                    f"{ex.nombre} {ex.apellido or ''}".strip())})
        else:
            nuevo = models.Cliente(
                nombre=cf.get("nombre") or "Garante",
                apellido=cf.get("apellido"),
                razon_social=cf.get("razon_social"),
                documento=cf.get("documento"),
                tipo_documento=cf.get("tipo_documento"),
                nacionalidad=cf.get("nacionalidad"),
                email=cf.get("email"),
                telefono=cf.get("telefono"),
                rol=models.ClienteRol.inquilino,
                notas="[IMPORTADO desde contrato — GARANTE / CO-FIRMANTE]",
                is_demo=is_demo,
            )
            db.add(nuevo); db.flush()
            cofirmantes_creados.append({"id": nuevo.id, "reutilizado": False,
                                         "nombre": f"{nuevo.nombre} {nuevo.apellido or ''}".strip()})
    resumen["co_firmantes"] = cofirmantes_creados

    # 3. Propiedad
    propd = datos.get("propiedad") or {}
    if not propd.get("direccion"):
        raise ValueError("Falta la dirección de la propiedad para poder importar.")
    propiedad = _buscar_propiedad_existente(db, propd, is_demo)
    if propiedad:
        resumen["propiedad"] = {"id": propiedad.id, "reutilizado": True}
        # Reasignar propietario si la propiedad no tenía uno y ahora sí
        if not propiedad.propietario_id:
            propiedad.propietario_id = propietario.id
        # Agregar los co-propietarios del PDF que NO estén ya vinculados.
        # No borramos los que ya tenía la propiedad reutilizada — el operador
        # puede limpiar manualmente si quiere reemplazar.
        ya_vinculados = {pp.cliente_id for pp in (propiedad.propietarios or [])}
        for i, (cli, porc, _reu) in enumerate(propietarios_creados):
            if cli.id in ya_vinculados:
                continue
            db.add(models.PropiedadPropietario(
                propiedad_id=propiedad.id,
                cliente_id=cli.id,
                porcentaje=porc,
                es_principal=False,  # ya hay un principal previo
            ))
        db.flush()
    else:
        tipo = propd.get("tipo") or "departamento"
        modalidad = propd.get("modalidad") or "alquiler"
        try:
            tipo_enum = models.PropiedadTipo(tipo)
        except ValueError:
            tipo_enum = models.PropiedadTipo.departamento
        try:
            mod_enum = models.PropiedadModalidad(modalidad)
        except ValueError:
            mod_enum = models.PropiedadModalidad.alquiler
        propiedad = models.Propiedad(
            direccion=propd["direccion"],
            ciudad=propd.get("ciudad"),
            provincia=propd.get("provincia"),
            tipo=tipo_enum,
            modalidad=mod_enum,
            estado=models.PropiedadEstado.ocupada if inquilino else models.PropiedadEstado.disponible,
            superficie_m2=propd.get("superficie_m2"),
            ambientes=propd.get("ambientes"),
            descripcion=(propd.get("descripcion") or "") + " [IMPORTADO desde PDF]",
            precio_alquiler=propd.get("precio_alquiler") or 0,
            expensas=propd.get("expensas") or 0,
            tasa_municipal=propd.get("tasa_municipal") or 0,
            propietario_id=propietario.id,
            is_demo=is_demo,
        )
        db.add(propiedad); db.flush()
        resumen["propiedad"] = {"id": propiedad.id, "reutilizado": False}
        # Crear filas pivote para cada co-propietario del PDF
        for i, (cli, porc, _reu) in enumerate(propietarios_creados):
            db.add(models.PropiedadPropietario(
                propiedad_id=propiedad.id,
                cliente_id=cli.id,
                porcentaje=porc,
                es_principal=(i == 0),
            ))
        db.flush()

    # 4. Contrato
    tipo_c = datos.get("tipo") or "alquiler_vivienda"
    estado_c = datos.get("estado") or "vigente"
    indice_c = datos.get("indice_ajuste") or "ipc"
    try:
        tipo_enum = models.ContratoTipo(tipo_c)
    except ValueError:
        tipo_enum = models.ContratoTipo.alquiler_vivienda
    try:
        estado_enum = models.ContratoEstado(estado_c)
    except ValueError:
        estado_enum = models.ContratoEstado.vigente
    try:
        indice_enum = models.IndiceAjuste(indice_c)
    except ValueError:
        indice_enum = models.IndiceAjuste.ipc

    from app.routers.contratos import _generar_codigo
    codigo = _generar_codigo(db, is_demo)

    from datetime import date as _date
    def _parse_date(s):
        if not s:
            return None
        try:
            return _date.fromisoformat(s)
        except Exception:
            return None

    # Determinar estado automáticamente según fechas si no vino explícito.
    fi = _parse_date(datos.get("fecha_inicio"))
    ff = _parse_date(datos.get("fecha_fin"))
    if not datos.get("estado"):
        from datetime import date as _date_today
        today = _date_today.today()
        if fi and ff and fi <= today <= ff:
            estado_enum = models.ContratoEstado.vigente
        elif ff and ff < _date_today.today():
            estado_enum = models.ContratoEstado.vencido
        else:
            estado_enum = models.ContratoEstado.borrador

    # Policies → JSON string
    policies_json = None
    if datos.get("policies"):
        policies_json = json.dumps(datos["policies"], ensure_ascii=False)

    contrato = models.Contrato(
        codigo=codigo,
        tipo=tipo_enum,
        estado=estado_enum,
        propiedad_id=propiedad.id,
        inquilino_id=inquilino.id if inquilino else None,
        fecha_inicio=fi,
        fecha_fin=ff,
        monto_inicial=datos.get("monto_inicial") or 0,
        deposito=datos.get("deposito") or 0,
        indice_ajuste=indice_enum,
        periodicidad_meses=datos.get("periodicidad_meses") or 6,
        porcentaje_fijo=datos.get("porcentaje_fijo") or 0,
        comision_porc=datos.get("comision_porc") or 0,
        mora_diaria_porc=datos.get("mora_diaria_porc") or 0,
        dia_inicio_pago=datos.get("dia_inicio_pago") or 1,
        dia_vencimiento_pago=datos.get("dia_vencimiento_pago") or 10,
        policies=policies_json,
        inventario=datos.get("inventario"),
        notas=(datos.get("notas") or "") + "\n[IMPORTADO desde contrato]",
        is_demo=is_demo,
    )
    db.add(contrato); db.flush()
    resumen["contrato"] = {"id": contrato.id, "codigo": contrato.codigo}

    # 4bis. Filas pivote contrato_inquilinos — una por cada firmante detectado.
    # El primero queda marcado como principal (compat con contrato.inquilino_id).
    for idx, (cli, _reu, _principal) in enumerate(inquilinos_creados):
        db.add(models.ContratoInquilino(
            contrato_id=contrato.id,
            cliente_id=cli.id,
            es_principal=(idx == 0),
            rol=("inquilino" if idx == 0 else "co_inquilino"),
        ))
    db.flush()

    # 5. Pagos futuros — sólo si el contrato está vigente y nos pidieron
    # autogenerar. Por defecto: sí (el operador puede borrar luego).
    pagos_generados = 0
    if (datos.get("generar_pagos_futuros", True) and
            estado_enum == models.ContratoEstado.vigente and fi and ff):
        pagos_generados = _generar_pagos_futuros(db, contrato)
    resumen["pagos_generados"] = pagos_generados

    db.commit()
    db.refresh(contrato)
    return resumen


def _generar_pagos_futuros(db: Session, contrato: models.Contrato) -> int:
    """Crea un Pago en estado=pendiente para cada mes desde hoy (o desde el
    inicio del contrato si todavía no empezó) hasta fecha_fin inclusive.

    Devuelve la cantidad de pagos creados.
    """
    from datetime import date as _date
    from dateutil.relativedelta import relativedelta

    if not contrato.fecha_inicio or not contrato.fecha_fin:
        return 0

    hoy = _date.today()
    # Arrancar desde el mes actual o el inicio del contrato (lo más reciente)
    cursor = max(contrato.fecha_inicio.replace(day=1), hoy.replace(day=1))
    end = contrato.fecha_fin

    dia_venc = contrato.dia_vencimiento_pago or 10
    creados = 0
    while cursor <= end:
        periodo = cursor.strftime("%Y-%m")
        # Evitar duplicar si ya existe el Pago de ese período
        existe = db.query(models.Pago).filter_by(
            contrato_id=contrato.id, periodo=periodo,
        ).first()
        if existe:
            cursor = cursor + relativedelta(months=1)
            continue
        # Fecha de vencimiento: día_venc del mes (clamp a último día si excede)
        try:
            fecha_venc = cursor.replace(day=dia_venc)
        except ValueError:
            # Día inválido (ej: día 31 en febrero) → último día del mes
            siguiente = (cursor.replace(day=1) + relativedelta(months=1))
            fecha_venc = siguiente - relativedelta(days=1)

        pago = models.Pago(
            contrato_id=contrato.id,
            periodo=periodo,
            fecha_vencimiento=fecha_venc,
            monto_total=contrato.monto_inicial,
            monto_alquiler=contrato.monto_inicial,
            estado=models.PagoEstado.pendiente,
            is_demo=contrato.is_demo,
        )
        db.add(pago)
        creados += 1
        cursor = cursor + relativedelta(months=1)

    db.flush()
    return creados
